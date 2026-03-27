"""
UsbExporter — экспорт tag_values в .xlsx и .docx на USB-флешку.
Вызывается автоматически при вставке флешки.
Во время записи периодически пищит 200 Гц / 100 мс.
"""
import subprocess
import threading
import time
import logging
from datetime import datetime
from pathlib import Path

from database import SessionLocal
from models import TagHistory

log = logging.getLogger(__name__)

# "idle" | "waiting" | "writing" | "done" | "error"
_status: str = "idle"


def get_status() -> str:
    return _status


def _set_status(s: str):
    global _status
    _status = s


def _get_partition(disk_node: str, timeout: float = 5.0) -> str | None:
    """Ждём появления первого раздела диска (например /dev/sdb → /dev/sdb1)."""
    import os
    disk_name = disk_node.replace("/dev/", "")
    deadline = time.time() + timeout
    while time.time() < deadline:
        sys_path = f"/sys/block/{disk_name}"
        try:
            for entry in os.listdir(sys_path):
                if entry.startswith(disk_name):
                    return f"/dev/{entry}"
        except OSError:
            pass
        time.sleep(0.3)
    return None


def _find_mount_point(device: str) -> str | None:
    """Проверить, смонтировано ли устройство. Вернуть точку монтирования или None."""
    try:
        with open("/proc/mounts") as f:
            for line in f:
                parts = line.split()
                if len(parts) >= 2 and parts[0] == device:
                    return parts[1]
    except OSError:
        pass
    return None


_MOUNT_POINT = "/mnt/usb"


def _mount_partition(partition: str) -> str | None:
    """Смонтировать раздел через sudo mount. Вернуть точку монтирования или None."""
    mount_point = _find_mount_point(partition)
    if mount_point:
        return mount_point
    try:
        import os
        os.makedirs(_MOUNT_POINT, exist_ok=True)
        # uid=1000 нужен для FAT/exFAT/NTFS чтобы user мог писать
        result = subprocess.run(
            ["sudo", "mount", "-o", "uid=1000,gid=1000,umask=002", partition, _MOUNT_POINT],
            capture_output=True, text=True, timeout=10,
        )
        if result.returncode != 0:
            # Для Linux-FS (ext4 и т.п.) опции uid/gid не поддерживаются — монтируем без них
            result = subprocess.run(
                ["sudo", "mount", partition, _MOUNT_POINT],
                capture_output=True, text=True, timeout=10,
            )
        if result.returncode == 0:
            return _MOUNT_POINT
        log.error("mount failed: %s", result.stderr.strip())
    except (FileNotFoundError, subprocess.TimeoutExpired) as e:
        log.error("mount exception: %s", e)
    return None


def _beep(freq: int, duration_ms: int):
    try:
        subprocess.Popen(
            ["beep", "-f", str(freq), "-l", str(duration_ms)],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
        )
    except FileNotFoundError:
        pass


def _beep_worker(stop_event: threading.Event):
    """Пищит 200 Гц / 100 мс непрерывно (каждые ~100 мс) пока stop_event не установлен."""
    while not stop_event.is_set():
        try:
            subprocess.Popen(
                ["beep", "-f", "200", "-l", "100"],
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
            ).wait()
        except FileNotFoundError:
            break


def _export_xlsx(path: Path, rows: list):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "История тегов"
    ws.append(["#", "Tag ID", "Название", "Значение", "Время"])
    for i, r in enumerate(rows, 1):
        ws.append([i, r.tag_id, r.tag_name, r.value, str(r.recorded_at)])
    wb.save(path)


def _export_docx(path: Path, rows: list):
    from docx import Document
    doc = Document()
    doc.add_heading("История тегов OPC UA", 0)
    doc.add_paragraph(f"Экспорт: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  Записей: {len(rows)}")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Tag ID"
    hdr[1].text = "Значение"
    hdr[2].text = "Время"
    hdr[3].text = "Название"
    for r in rows:
        cells = table.add_row().cells
        cells[0].text = r.tag_id
        cells[1].text = r.value
        cells[2].text = str(r.recorded_at)
        cells[3].text = r.tag_name
    doc.save(path)


def _do_export(device_info: dict):
    node = device_info.get("node", "")
    log.info("USB inserted: looking for partition on %s...", node)
    _set_status("waiting")

    partition = _get_partition(node)
    if not partition:
        log.error("No partition found on %s — export skipped", node)
        _set_status("error")
        return

    log.info("Found partition: %s, mounting...", partition)
    mount_point = _mount_partition(partition)
    if not mount_point:
        log.error("Could not mount %s — export skipped", partition)
        _set_status("error")
        return

    log.info("Exporting to %s", mount_point)
    _set_status("writing")

    stop_beep = threading.Event()
    beep_thread = threading.Thread(target=_beep_worker, args=(stop_beep,), daemon=True)
    beep_thread.start()

    try:
        db = SessionLocal()
        try:
            rows = db.query(TagHistory).order_by(TagHistory.recorded_at).all()
        finally:
            db.close()

        ts = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        dest = Path(mount_point)

        xlsx_path = dest / f"tags_{ts}.xlsx"
        docx_path = dest / f"tags_{ts}.docx"

        _export_xlsx(xlsx_path, rows)
        xlsx_size = xlsx_path.stat().st_size if xlsx_path.exists() else 0
        if xlsx_size == 0:
            raise RuntimeError(f"XLSX file missing or empty after write: {xlsx_path}")
        log.info("XLSX written: %s (%d rows, %d bytes)", xlsx_path.name, len(rows), xlsx_size)

        _export_docx(docx_path, rows)
        docx_size = docx_path.stat().st_size if docx_path.exists() else 0
        if docx_size == 0:
            raise RuntimeError(f"DOCX file missing or empty after write: {docx_path}")
        log.info("DOCX written: %s (%d rows, %d bytes)", docx_path.name, len(rows), docx_size)

        _set_status("done")
    except Exception:
        log.exception("USB export failed")
        _set_status("error")
    finally:
        stop_beep.set()
        beep_thread.join(timeout=1.0)
        if _status == "done":
            _beep(1000, 400)
        log.info("USB export finished")


def export_on_insert(device_info: dict):
    """Точка входа — запускает экспорт в фоновом потоке."""
    threading.Thread(
        target=_do_export,
        args=(device_info,),
        daemon=True,
        name="usb-export",
    ).start()
