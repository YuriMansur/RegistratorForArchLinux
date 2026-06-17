"""
SessionExporter — экспорт данных сессии в xlsx и docx на диск сервера.
Вызывается когда End=True.
"""
import logging
from datetime import datetime, timezone
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import matplotlib.ticker as mticker

# signals — маппинг технического имени тега в подпись и единицу измерения.
from services import signals

EXPORT_DIR = Path("/home/user/registrator/exports")

# Теги, которые не попадают в экспорт (управляющие, служебные и т.д.)
# Добавляй сюда имена тегов (Tag.name) которые не нужны в таблице.
_EXCLUDE_TAG_NAMES: set[str] = {
    "inProcess",
    "End",
}

log = logging.getLogger(__name__)


def _to_local(dt: datetime) -> str:
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone().strftime("%Y-%m-%d %H:%M:%S")


def _fmt(dt: datetime) -> str:
    return dt.astimezone().strftime("%Y-%m-%d_%H-%M-%S")


def _fmt_duration(start: datetime, end: datetime) -> str:
    """Длительность испытания в виде «Ч ч ММ мин СС с» (реальное время часы/мин/сек)."""
    # Нормализуем таймзоны, чтобы вычесть без ошибки naive/aware.
    if start.tzinfo is None:
        start = start.replace(tzinfo=timezone.utc)
    if end.tzinfo is None:
        end = end.replace(tzinfo=timezone.utc)
    total = int((end - start).total_seconds())
    if total < 0:
        total = 0
    h, rem = divmod(total, 3600)
    m, s = divmod(rem, 60)
    return f"{h} ч {m:02d} мин {s:02d} с"


def export_session(session_start: datetime, session_end: datetime, test_id: int) -> None:
    """Экспортировать данные испытания (автоматический вызов при End=True)."""
    _export(test_id, session_start, session_end)


def export_by_test_id(test_id: int) -> None:
    """Экспортировать данные испытания по его ID (ручной вызов из GUI)."""
    from db.database import SessionLocal
    from db.models import Checkout

    db = SessionLocal()
    try:
        checkout = db.get(Checkout, test_id)
        if checkout is None:
            log.warning("export_by_test_id: checkout %s not found", test_id)
            return
        started_at = checkout.started_at.replace(tzinfo=timezone.utc) \
            if checkout.started_at.tzinfo is None else checkout.started_at
        ended_at   = checkout.ended_at or datetime.now(timezone.utc)
        if ended_at.tzinfo is None:
            ended_at = ended_at.replace(tzinfo=timezone.utc)
    finally:
        db.close()

    _export(test_id, started_at, ended_at)


def export_by_date_range(from_dt: datetime, to_dt: datetime) -> None:
    """Экспортировать данные по произвольному диапазону дат (ручной вызов из GUI)."""
    from db.database import SessionLocal
    from db.models import TagHistory, Tag

    if from_dt.tzinfo is None:
        from_dt = from_dt.replace(tzinfo=timezone.utc)
    if to_dt.tzinfo is None:
        to_dt = to_dt.replace(tzinfo=timezone.utc)

    from_naive = from_dt.replace(tzinfo=None)
    to_naive   = to_dt.replace(tzinfo=None)

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)

    db = SessionLocal()
    try:
        rows = (
            db.query(TagHistory, Tag)
            .outerjoin(Tag, TagHistory.tag_id == Tag.id)
            .filter(TagHistory.recorded_at >= from_naive)
            .filter(TagHistory.recorded_at <= to_naive)
            .order_by(TagHistory.recorded_at)
            .all()
        )
    finally:
        db.close()

    if not rows:
        log.warning("Export range: no rows found for %s — %s", from_dt, to_dt)
        return

    folder_name = f"Data_{_fmt(from_dt)}_{_fmt(to_dt)}"
    session_dir = EXPORT_DIR / folder_name
    session_dir.mkdir(parents=True, exist_ok=True)

    try:
        ts = _fmt(to_dt)
        headers, data, header_base = _pivot(rows)
        _write_xlsx(session_dir / f"data_{ts}.xlsx", headers, data)
        pngs = _write_png_per_tag(session_dir, headers, data,
                                  only_headers=_chart_headers(header_base))
        _write_docx(session_dir / f"data_{ts}.docx", headers, data,
                    title=f"Данные {_to_local(from_dt)} — {_to_local(to_dt)}")
        _write_docx_charts(session_dir / f"data_{ts}_charts.docx", pngs,
                           title=f"Данные {_to_local(from_dt)} — {_to_local(to_dt)} — графики",
                           duration=_fmt_duration(from_dt, to_dt))
        log.info("Export range done: %d rows → %s", len(rows), session_dir.name)
    except Exception:
        log.exception("Export range failed")


def _export(test_id: int, session_start: datetime, session_end: datetime) -> None:
    """Общая логика: запросить TagHistory по test_id, записать xlsx и docx."""
    from db.database import SessionLocal
    from db.models import TagHistory, Tag

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)

    db = SessionLocal()
    try:
        rows = (
            db.query(TagHistory, Tag)
            .outerjoin(Tag, TagHistory.tag_id == Tag.id)
            .filter(TagHistory.test_id == test_id)
            .order_by(TagHistory.recorded_at)
            .all()
        )
    finally:
        db.close()

    if not rows:
        log.warning("Export: no rows found for test_id=%s", test_id)
        return

    session_dir = EXPORT_DIR / f"checkout_{test_id}_{_fmt(session_start)}_{_fmt(session_end)}"
    session_dir.mkdir(parents=True, exist_ok=True)

    ts = _fmt(session_end)
    try:
        headers, data, header_base = _pivot(rows)
        _write_xlsx(session_dir / f"session_{ts}.xlsx", headers, data)
        pngs = _write_png_per_tag(session_dir, headers, data,
                                  only_headers=_chart_headers(header_base))
        _write_docx(session_dir / f"session_{ts}.docx", headers, data, title=f"Испытание №{test_id}")
        _write_docx_charts(session_dir / f"session_{ts}_charts.docx", pngs,
                           title=f"Испытание №{test_id} — графики",
                           duration=_fmt_duration(session_start, session_end))
        log.info("Export done: %d rows → %s", len(rows), session_dir.name)
    except Exception:
        log.exception("Export failed for test_id=%s", test_id)


def _pivot(rows: list):
    """Сгруппировать строки по времени — каждый тег становится колонкой.
    Возвращает:
        headers      — список заголовков колонок: "Имя [единицы]"
        data         — dict {recorded_at: {header: value}}
        header_base  — dict {header: базовое имя тега} (для фильтрации графиков)
    """
    from collections import defaultdict
    # Сохраняем порядок тегов по первому появлению
    headers = []
    data = defaultdict(dict)
    # Карта "заголовок колонки → базовое имя тега" — нужна чтобы отобрать
    # графики по chart_tags из конфига (там заданы именно базовые имена тегов).
    header_base: dict[str, str] = {}
    for h, tag in rows:
        name  = tag.name  if tag else str(h.tag_id)
        # Технические управляющие теги исключаем по короткому имени до перевода в подпись.
        # Для массивов вида "inProcess[0]" сравниваем по базовому имени.
        base = name.split("[", 1)[0]
        if base in _EXCLUDE_TAG_NAMES:
            continue
        # Подпись и единица — из signals.json (фоллбек на техническое имя, если тег не описан).
        label = signals.get_label(name)
        units = signals.get_unit(name)
        header = f"{label} [{units}]" if units else label
        if header not in headers:
            headers.append(header)
            header_base[header] = base
        try:
            data[h.recorded_at][header] = f"{float(h.value):.2f}"
        except (ValueError, TypeError):
            data[h.recorded_at][header] = h.value if h.value is not None else ""
    return headers, data, header_base


def _write_xlsx(path: Path, headers: list, data: dict) -> None:
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Сессия"

    ws.append(["Время"] + headers)

    for ts in sorted(data.keys()):
        row = [_to_local(ts)]
        for header in headers:
            row.append(data[ts].get(header, ""))
        ws.append(row)

    wb.save(path)


def _write_docx(path: Path, headers: list, data: dict, title: str = "") -> None:
    from docx import Document
    from docx.shared import Mm, Pt

    doc = Document()

    section = doc.sections[0]
    section.orientation = 1
    section.page_width  = Mm(420)
    section.page_height = Mm(297)
    section.left_margin   = Mm(10)
    section.right_margin  = Mm(10)
    section.top_margin    = Mm(10)
    section.bottom_margin = Mm(10)

    if title:
        doc.add_heading(title, 0)

    timestamps = sorted(data.keys())

    table = doc.add_table(rows=1 + len(timestamps), cols=1 + len(headers))
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Время"
    for i, header in enumerate(headers, 1):
        hdr[i].text = header

    for row_idx, ts in enumerate(timestamps, 1):
        cells = table.rows[row_idx].cells
        cells[0].text = _to_local(ts)
        for i, header in enumerate(headers, 1):
            cells[i].text = data[ts].get(header, "")

    doc.save(path)


def _add_protocol_header(doc, duration: str = "") -> None:
    """Вставить в начало документа шапку «Протокол испытания» (форма ПКБА).
    Значения пустые (заполняются вручную в Word), кроме «Время испытания» —
    туда подставляется реальная длительность испытания (часы/мин/сек)."""
    from docx.shared import Pt, Mm, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Палитра: тёмно-синяя плашка-шапка, светло-синие подписи, белые поля.
    TITLE_BG  = "1F4E79"   # фон строки с названием организации
    LABEL_BG  = "D9E2F3"   # фон ячеек-подписей
    WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
    ACCENT    = RGBColor(0x1F, 0x4E, 0x79)  # цвет значения «Время испытания»
    _ALIGN = {"left": WD_ALIGN_PARAGRAPH.LEFT,
              "center": WD_ALIGN_PARAGRAPH.CENTER,
              "right": WD_ALIGN_PARAGRAPH.RIGHT}

    def shade(cell, fill_hex):
        """Залить фон ячейки цветом (через w:shd в свойствах ячейки)."""
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tcPr.append(shd)

    def fill(cell, text, bold=False, size=8, align="left", color=None, bg=None):
        """Записать текст в ячейку (многострочный — через \\n): жирность, размер,
        выравнивание, цвет текста и фон. Текст по вертикали центрируется."""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if bg:
            shade(cell, bg)
        cell.text = ""
        first = cell.paragraphs[0]
        for i, line in enumerate(str(text).split("\n")):
            p = first if i == 0 else cell.add_paragraph()
            p.alignment = _ALIGN.get(align, WD_ALIGN_PARAGRAPH.LEFT)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after  = Pt(0)
            run = p.add_run(line)
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = "Calibri"
            if color is not None:
                run.font.color.rgb = color

    def label(cell, text):
        """Ячейка-подпись: жирная, на светло-синем фоне."""
        fill(cell, text, bold=True, bg=LABEL_BG)

    # 6 колонок: три пары «подпись : значение». 11 строк: организация, блок полей, низ.
    table = doc.add_table(rows=11, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    # ── Строка 0: заголовок протокола на всю ширину (название организации убрано) ─
    fill(table.cell(0, 0).merge(table.cell(0, 5)),
         "Протокол испытания\n№ ____________        Дата ____________",
         bold=True, size=10, align="center", color=WHITE, bg=TITLE_BG)

    # ── Строки 1–6: три колонки подписей (значения пустые) ───────────────────────
    left  = ["Предприятие", "Заказчик", "Номер заказа", "МСЛ", "Состав", "Исполнитель"]
    mid   = ["Вид арматуры", "Обозначение", "Зав №", "Производитель", "Уплотнение"]
    right = ["DN, мм", "PN", "t вод, С", "t возд, С"]
    for i, text in enumerate(left):
        label(table.cell(1 + i, 0), text)
    for i, text in enumerate(mid):
        label(table.cell(1 + i, 2), text)
    for i, text in enumerate(right):
        label(table.cell(1 + i, 4), text)

    # ── Строка 7: «Испытание на прочность» + широкое свободное поле ───────────────
    label(table.cell(7, 0), "Испытание на прочность")
    table.cell(7, 1).merge(table.cell(7, 5))

    # ── Строки 8–10: давления / среда / результат / время ────────────────────────
    label(table.cell(8, 0), "Начальное давление")
    label(table.cell(8, 2), "Испытательная среда")
    table.cell(8, 3).merge(table.cell(8, 5))
    label(table.cell(9, 0), "Минимальное давление")
    label(table.cell(9, 2), "Результат")
    table.cell(9, 3).merge(table.cell(9, 5))
    label(table.cell(10, 0), "Время испытания")
    # В объединённую ячейку справа — реальная длительность, выделена жирным и цветом.
    fill(table.cell(10, 2).merge(table.cell(10, 5)), duration,
         bold=True, size=10, color=ACCENT)

    # ── Ширины колонок (фиксируем, чтобы таблица не «плясала» по содержимому) ─────
    # 3 пары подпись/значение; сумма ≈ ширине рабочей области A4 книжной (194 мм).
    widths = [Mm(30), Mm(34), Mm(30), Mm(34), Mm(28), Mm(38)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]


def _write_docx_charts(path: Path, images: list[Path] | None, title: str = "",
                       duration: str = "") -> None:
    """Отдельный docx только с графиками (PNG), формат A4 КНИЖНЫЙ.
    Графики идут подряд с мелким отступом друг от друга (без разрывов страниц):
    Word сам переносит на новую страницу, когда текущая заполнилась."""
    if not images:
        return

    from docx import Document
    from docx.shared import Mm, Pt
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # A4 книжный: ширина 210 < высота 297 мм (высота > ширины = книжная).
    # Поля минимальные — чтобы на странице помещалось больше графиков.
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width  = Mm(210)
    section.page_height = Mm(297)
    section.left_margin   = Mm(8)
    section.right_margin  = Mm(8)
    section.top_margin    = Mm(8)
    section.bottom_margin = Mm(8)

    # Шапка протокола испытания (форма ПКБА) — поля под ручное заполнение,
    # кроме «Время испытания» (туда подставляется реальная длительность).
    _add_protocol_header(doc, duration=duration)
    # Небольшой отступ между шапкой и первым графиком.
    doc.add_paragraph()

    # Ширина рабочей области = ширина листа минус поля.
    usable_width = section.page_width - section.left_margin - section.right_margin
    for img in images:
        # Пропускаем отсутствующие файлы — одна битая картинка не валит весь docx.
        if not Path(img).exists():
            continue
        # Добавляем картинку (в свой параграф); разрывов страниц НЕ ставим —
        # графики идут впритык, перенос на новую страницу делает сам Word.
        doc.add_picture(str(img), width=usable_width)
        # Мелкий зазор между графиками: сверху 0, снизу 6 pt (~2 мм) — не впритык, но плотно.
        p = doc.paragraphs[-1]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(6)
        pf.line_spacing = 1.0

    doc.save(path)


def _chart_headers(header_base: dict[str, str]) -> set[str] | None:
    """Множество заголовков колонок, чьи графики должны попасть в *_charts.docx —
    по списку chart_tags из конфига (там заданы базовые имена тегов).
    Возвращает None, если фильтр не задан (в docx идут графики всех тегов)."""
    allowed = signals.get_chart_tags()
    if not allowed:
        return None
    return {h for h, base in header_base.items() if base in allowed}


def _write_png_per_tag(session_dir: Path, headers: list, data: dict,
                       only_headers: set[str] | None = None) -> list[Path]:
    """Нарисовать по графику на тег и сохранить trend_*.png.
    only_headers — если задано, рисуем только эти заголовки (фильтр chart_tags из конфига);
    None = рисуем графики всех тегов.
    Возвращает список путей созданных PNG (в порядке headers) — нужен чтобы
    встроить эти же картинки в docx (см. _write_docx)."""
    if not headers or not data:
        return []

    from collections import defaultdict
    local_tz = datetime.now(timezone.utc).astimezone().tzinfo
    timestamps = sorted(data.keys())

    # Пути созданных PNG — накапливаем для встраивания в Word.
    created: list[Path] = []

    for header in headers:
        # Фильтр из конфига: тег не в chart_tags → график не строим.
        if only_headers is not None and header not in only_headers:
            continue
        tag_times = []
        values = []
        for ts in timestamps:
            v = data[ts].get(header, None)
            if v is None:
                continue
            try:
                val = float(v)
            except (ValueError, TypeError):
                continue
            dt = ts.replace(tzinfo=timezone.utc) if ts.tzinfo is None else ts
            tag_times.append(dt.astimezone(local_tz))
            values.append(val)

        if len(tag_times) < 2:
            continue

        # Усреднение по 1-минутным интервалам
        buckets: dict = defaultdict(list)
        for dt, val in zip(tag_times, values):
            minute = dt.replace(second=0, microsecond=0)
            buckets[minute].append(val)
        avg_times = sorted(buckets)
        avg_values = [sum(buckets[m]) / len(buckets[m]) for m in avg_times]

        if len(avg_times) < 2:
            continue

        DPI = 150
        width = max(24, min(len(avg_times), 32767 // DPI))
        fig, ax = plt.subplots(figsize=(width, 6), dpi=DPI)
        ax.plot(avg_times, avg_values, linewidth=1.2, color="steelblue",
                marker="o", markersize=3)

        # ── Отметить максимум: точка на кривой + подпись значения (до десятых) ──
        max_idx = max(range(len(avg_values)), key=lambda i: avg_values[i])
        max_x, max_y = avg_times[max_idx], avg_values[max_idx]
        ax.plot(max_x, max_y, marker="o", markersize=8, color="crimson", zorder=5)
        ax.annotate(
            f"{max_y:.1f}",
            xy=(max_x, max_y), xytext=(0, 10), textcoords="offset points",
            ha="center", va="bottom", fontsize=11, fontweight="bold", color="crimson",
            bbox=dict(boxstyle="round,pad=0.25", fc="white", ec="crimson", alpha=0.85),
            zorder=6,
        )
        # Запас сверху по Y, чтобы подпись максимума не обрезалась краем графика.
        y_min, y_max = min(avg_values), max(avg_values)
        y_range = (y_max - y_min) or 1.0
        ax.set_ylim(y_min - y_range * 0.05, y_max + y_range * 0.15)

        ax.set_xlim(avg_times[0], avg_times[-1])
        # Подписи оси Y — до десятых.
        ax.yaxis.set_major_formatter(mticker.FormatStrFormatter("%.1f"))
        ax.xaxis.set_major_formatter(mdates.DateFormatter("%H:%M", tz=local_tz))
        ax.xaxis.set_major_locator(mdates.AutoDateLocator())
        fig.autofmt_xdate()
        ax.set_xlabel("Время")
        ax.set_ylabel(header)
        t_start = avg_times[0].strftime("%d.%m.%Y %H:%M")
        t_end   = avg_times[-1].strftime("%d.%m.%Y %H:%M")
        ax.set_title(f"{header}\n{t_start} — {t_end}", fontsize=11)
        ax.grid(True, linestyle="--", alpha=0.5)
        fig.tight_layout()

        safe_name = header.replace("/", "_").replace(" ", "_").replace("[", "").replace("]", "")
        out_path = session_dir / f"trend_{safe_name}.png"
        fig.savefig(out_path, dpi=150)
        plt.close(fig)
        created.append(out_path)

    return created
